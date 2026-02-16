from io import BytesIO
from pathlib import Path
import os
import tempfile
import csv
import pandas as pd
from pandas.errors import EmptyDataError, ParserError
from flask import Flask, request, jsonify, render_template, send_file, make_response

import uuid


from docx import Document
from docx.shared import Pt


ALLOWED_COLUMNS = [
    "document_type",
    "document_number",
    "document_date",
    "supplier_name",
    "customer_name",
    "shipment_location",
    "payment_terms",
    "line_number",
    "item_code",
    "item_description",
    "category",
    "quantity",
    "unit_of_measure",
    "unit_price",
    "extended_price"
]


from demo import (
    extract_structured_data_from_tabular_with_llm,
    extract_text,
    extract_structured_data_with_llm,
    json_to_rows,
    load_config,
    get_openai_client,
    read_tabular,
    read_tabular_file_to_rows,
)

UPLOAD_DIR = Path("uploads")
DATA_DIR = Path("data")
CSV_PATH = DATA_DIR / "dataset.csv"

SUPPORTED_SUFFIXES = {".docx", ".pdf", ".jpg", ".jpeg", ".png", ".csv", ".xlsx"}

# ---- Per-browser "database" via persistent cookie ----

SESSION_DATA_DIR = DATA_DIR / "sessions"
SESSION_DATA_DIR.mkdir(exist_ok=True)

DB_ID_COOKIE = "db_id"
PERSIST_DAYS = 7  # change to 365 if you want ~1 year persistence

def get_db_csv_path(db_id: str) -> Path:
    # A separate CSV per browser
    return SESSION_DATA_DIR / f"{db_id}.csv"

def get_or_create_db_id() -> str:
    db_id = request.cookies.get(DB_ID_COOKIE)
    if db_id and isinstance(db_id, str) and len(db_id) >= 8:
        return db_id
    return str(uuid.uuid4())

def get_active_csv_path() -> Path:
    db_id = get_or_create_db_id()
    return get_db_csv_path(db_id)

def set_persistent_db_cookie(resp, db_id: str):
    resp.set_cookie(
        DB_ID_COOKIE,
        db_id,
        max_age=PERSIST_DAYS * 24 * 3600,
        samesite="Lax",
        httponly=True,
    )
    return resp


app = Flask(__name__)

# Make sure folders exist
UPLOAD_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)

# Init OpenAI client
config = load_config()
client = get_openai_client(config["api_key"])
model_name = config["model"]

def set_table_font_size(table, size=5):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)

def fill_missing_by_commas(filepath):
    new_lines = []

    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip("\n")  # remove newline temporarily

            # Replace consecutive commas: ",," → ",Not Defined,"
            while ",," in line:
                line = line.replace(",,", ",Not Defined,")

            # If line ends with a comma → add "Not Defined"
            if line.endswith(","):
                line += "Not Defined"

            new_lines.append(line)

    # Write back to the same file
    with open(filepath, "w", encoding="utf-8") as f:
        for line in new_lines:
            f.write(line + "\n")

 


def safe_read_csv(csv_path: Path) -> pd.DataFrame:
    """
    Read the CSV used as our database.

    - If it does not exist or is zero bytes, return an empty DataFrame.
    - If it is malformed, rename it to .bak and return an empty DataFrame.
    """
    if not csv_path.exists() or csv_path.stat().st_size == 0:
        return pd.DataFrame()

    try:
        fill_missing_by_commas(csv_path)
        return pd.read_csv(csv_path)
    except (EmptyDataError, ParserError):
        backup = csv_path.with_suffix(csv_path.suffix + ".bak")
        try:
            csv_path.rename(backup)
        except OSError:
            pass
        return pd.DataFrame()
    except Exception as e:
        app.logger.exception("Unexpected error reading CSV %s: %s", csv_path, e)
        return pd.DataFrame()

def append_rows_to_csv(rows, csv_path: Path):
    """Append rows (list of dict) to CSV, creating it if needed."""
    if not rows:
        return

    new_df = pd.DataFrame(rows)

    existing = safe_read_csv(csv_path)
    combined = pd.concat([existing, new_df], ignore_index=True)

    combined.to_csv(csv_path, index=False)
    fill_missing_by_commas(csv_path)
    app.logger.info(
        "CSV updated with %d new row(s). Total rows now: %d",
        len(new_df),
        len(combined),
    )

def reset_csv(csv_path: Path):
 
    existing = safe_read_csv(csv_path)
    existing = existing.iloc[0:0]
    existing.to_csv(csv_path, index=False)

    app.logger.info(
        "CSV updated with 0 new row(s). Total rows now: %d",
        len(existing)
    )


def safe_read_csv(csv_path: Path) -> pd.DataFrame:
    """
    Read the CSV used as our database.

    - If it does not exist or is zero bytes, return an empty DataFrame.
    - If it is malformed, rename it to a .bak file and return an empty DataFrame.
    - Never raise to the caller.
    """
    if not csv_path.exists() or csv_path.stat().st_size == 0:
        return pd.DataFrame()

    try:
        return pd.read_csv(csv_path)
    except (EmptyDataError, ParserError):
        # Corrupted CSV, keep a backup and start fresh
        backup = csv_path.with_suffix(csv_path.suffix + ".bak")
        try:
            csv_path.rename(backup)
        except OSError:
            # If rename fails, just ignore and treat it as empty
            pass
        return pd.DataFrame()
    except Exception as e:
        app.logger.exception("Unexpected error reading CSV %s: %s", csv_path, e)
        return pd.DataFrame()


@app.route("/")
def index():
    # Renders templates/index.html
    db_id = get_or_create_db_id()
    resp = make_response(render_template("index.html"))
    set_persistent_db_cookie(resp, db_id)
    return resp


@app.route("/data", methods=["GET"])
def get_data():
    """Return current CSV dataset as JSON."""
    try:
        df = safe_read_csv(get_active_csv_path())

        # If no data or no columns, send an empty list
        if df.empty or df.shape[1] == 0:
            return jsonify([])

        # Replace NaN with None so JSON is valid (null)
       # df = df.where(pd.notnull(df), None)
   #     df = df.fillna(value=None)

        data = df.to_dict(orient="records")
        return jsonify(data)
    except Exception as e:
        app.logger.exception("Unexpected error in /data: %s", e)
        return jsonify([]), 500

@app.route("/download", methods=["GET"])
def download():
    #To get the filetype
    print("\n \n SOMEHTHING \n \n")
    file_type = request.args.get("type")  # e.g. "csv", "pdf"
    print(file_type)

    database_path  = get_active_csv_path() #It used to be './data/dataset.csv',

    if (file_type == 'csv' ) or (file_type == 'xlsx' )  :
        return send_file(database_path, as_attachment=True)
    
    df = pd.read_csv(database_path, index_col=None)

    if (file_type == 'docx'):
        doc = Document() #Create Word Document instance
        doc.add_heading("Document Processor - Database")
        doc.add_heading("Made by the AI4ID Consortium - Texas A&M")

        table = doc.add_table(rows=1 , cols=len(df.columns))
        hdr_cells = table.rows[0].cells #Cell Selector, positioning in Cell 0,0

        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
        set_table_font_size(table)
        file_io = BytesIO()
        doc.save(file_io)
        file_io.seek(0)
        return send_file(file_io, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name="dataset.docx")

@app.route("/reset", methods=["POST"])
def reset():
    print("Another Somethinf")
        #return jsonify({"rows": rows})
    csv_path = get_active_csv_path()
    reset_csv(csv_path)
   
    db_id = get_or_create_db_id()
    resp = make_response(jsonify({"rows": "[]" }))
    set_persistent_db_cookie(resp, db_id)
    return resp



@app.route("/upload", methods=["POST"])
def upload():
    """Receive a file, process it with demo.py, update CSV, return extracted rows."""
    if "file" not in request.files:
        return jsonify({"error": "No file in request"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    suffix = Path(file.filename).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        return jsonify({"error": f"Unsupported file type: {suffix}"}), 400

    # Save file to a temp path
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=UPLOAD_DIR) as tmp:
        tmp_path = Path(tmp.name)
        file.save(tmp_path)

    try:

        if suffix in ['.csv', '.xlsx','csv', 'xlsx']:
            info = read_tabular(str(tmp_path))
            structured = extract_structured_data_from_tabular_with_llm(client, model_name, info)
        else:
            raw_text = extract_text(str(tmp_path))
            structured = extract_structured_data_with_llm(client, model_name, raw_text)
        if not raw_text.strip():
            return jsonify({"error": "No text could be extracted from this document"}), 400
        
        rows = json_to_rows(structured)

        #append_rows_to_csv(rows, CSV_PATH)

        #return jsonify({"rows": rows})
        csv_path = get_active_csv_path()
        append_rows_to_csv(rows, csv_path)

        db_id = get_or_create_db_id()
        resp = make_response(jsonify({"rows": rows}))
        set_persistent_db_cookie(resp, db_id)
        return resp


    except Exception as e:
        # In a more serious app, do better error logging
        return jsonify({"error": str(e)}), 500
    finally:
        # Clean up the uploaded file
        if tmp_path.exists():
            tmp_path.unlink()


#if __name__ == "__main__":
 #    For local testing
    #fill_missing_by_commas(CSV_PATH)
    #app.run(host="0.0.0.0", port=5500, debug=True)

if __name__ == "__main__":
    app.run(debug=True)
